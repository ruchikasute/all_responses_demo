import streamlit as st
import os
import time
import re
from io import BytesIO
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
from docx import Document
from openai import AzureOpenAI
from langchain_openai import AzureOpenAIEmbeddings
from langchain_core.documents import Document as LDocument
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from Modules.prompts import (
    get_executive_summary_and_objective_prompt,
    get_scope_prereq_assumptions_prompt,
    get_resource_schedule_and_commercial_prompt,
    get_communication_plan_prompt
)



# -------------------------------------------------------
# 1. SETUP
# -------------------------------------------------------
load_dotenv()
KNOWLEDGE_FOLDER = "Knowledge_Repo"
PERSIST_DIR = "chroma_db"

st.set_page_config(page_title="RFP Proposal AI Generator", layout="wide")

# Custom CSS for Professional Look (Mimics AutoRFP style)
st.markdown("""
<style>
/* Primary Brand Color */
:root {
    --primary-blue: #1A75E0;
    --light-blue-bg: #EAF3FF;
}

/* Centering and large text for the main header */
.main-header {
    text-align: center;
    color: #000;
    font-size: 3em;
    font-weight: 800;
    padding-top: 20px;
    padding-bottom: 5px;
}
.highlight-text {
    color: var(--primary-blue);
}
.sub-tagline {
    text-align: center;
    color: #555;
    font-size: 1.1em;
    padding-bottom: 40px;
}

/* Style for the upload boxes (the two blocks requested) */
.upload-card {
    border: 1px solid #E0E0E0;
    border-radius: 12px;
    padding: 30px 20px;
    text-align: center;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    transition: all 0.2s ease-in-out;
    background-color: #F9F9F9;
    height: 100%;
    margin-bottom: 20px;
}
.upload-card:hover {
    box-shadow: 0 8px 16px rgba(0, 0, 0, 0.1);
    border-color: var(--primary-blue);
}
.upload-header {
    color: var(--primary-blue);
    font-weight: 600;
    margin-bottom: 10px;
}

/* Streamlit component tweaks */
div.stButton > button {
    background-color: var(--primary-blue);
    color: white;
    border-radius: 8px;
    border: none;
    padding: 10px 20px;
    font-weight: bold;
    transition: background-color 0.2s;
}
div.stButton > button:hover {
    background-color: #145CB0;
}
/* Style status box titles for consistency */
.stStatus [data-testid="stStatusContainer"] > div:first-child > div:first-child {
    font-weight: 600;
    color: #333;
}
</style>
""", unsafe_allow_html=True)


# -------------------------------------------------------
# 2. UTILITIES
# -------------------------------------------------------

def extract_text(file):
    """Extract text from PDF or DOCX"""
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""


from pinecone import Pinecone, ServerlessSpec
from langchain_pinecone import PineconeVectorStore
from langchain_community.embeddings import HuggingFaceEmbeddings
import os
@st.cache_resource(show_spinner=False)
def build_knowledge_base(folder="Knowledge_Repo"):
    embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index_name = "response-generator"

    # Create index if it doesn't exist
    if index_name not in [idx["name"] for idx in pc.list_indexes()]:
        pc.create_index(
            name=index_name,
            dimension=384,  # ✅ MiniLM-L6-v2 has 384 dims (not 1024)
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1")
        )

    index = pc.Index(index_name)

    vector_store = PineconeVectorStore(index=index, embedding=embedding_model)

    # --- Upload documents if index is empty ---
    stats = pc.describe_index(index_name)
    if stats.get("status", {}).get("ready", False):
        # Load local RFP references
        docs = []
        for file in os.listdir(folder):
            if file.endswith((".pdf", ".docx")):
                path = os.path.join(folder, file)
                text = extract_text(open(path, "rb"))
                if text.strip():
                    docs.append(LDocument(page_content=text, metadata={"source": file}))

        # Before uploading, check which files already exist
        existing_docs = [
            m.metadata.get("source") 
            for m in vector_store.similarity_search("test", k=10)
            if m.metadata and "source" in m.metadata
        ]

        new_docs = [d for d in docs if d.metadata["source"] not in existing_docs]

        if new_docs:
            vector_store.add_documents(new_docs)
            print(f"✅ Added {len(new_docs)} new docs to Pinecone")
        else:
            print("✅ Knowledge base already up to date — no new uploads")


    return vector_store

def condense_rfp_text(rfp_text, client):
    """
    Summarizes raw RFP text into a compact, structured form
    that retains all important details but reduces token size.
    """
    if len(rfp_text) < 1000:
        return rfp_text  # No need to condense small files


    prompt = f"""
    Summarize the following RFP content briefly.
    Keep only the important project details, scope, and objectives.
    Avoid unnecessary legal or boilerplate sections.

    RFP Content:
    {rfp_text[:5000]}
    """
    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.2,
        max_tokens=600,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# 🧠 Add this helper just below condense_rfp_text()
def get_most_relevant_reference(reference_text, condensed_rfp, keyword_hint=None):
    """
    Filters the large reference document to keep only the 2 most relevant paragraphs
    based on keyword overlap with the condensed RFP text.
    This reduces token size and speeds up prompt processing.
    """
    paragraphs = reference_text.split("\n\n")
    condensed_keywords = set(condensed_rfp.lower().split())

    if keyword_hint:
        condensed_keywords |= set(keyword_hint.lower().split())

    scores = []
    for p in paragraphs:
        overlap = len(set(p.lower().split()) & condensed_keywords)
        scores.append((overlap, p))
    best = sorted(scores, reverse=True)[:2]
    return "\n\n".join([p for _, p in best])

def apply_bullet_to_para(paragraph, list_id='1'):
    """
    Applies a dot bullet style (list level 0) using its XML structure.
    Uses numId='1' which is often the default bullet style in templates.
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    
    # Set the list level (0 is the main level)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    
    # Set the list ID (Most default templates use ID '1' for the first bullet definition)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), list_id)
    
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def insert_executive_summary_into_template(
    template_path,
    summary_text,
    objective_text=None,
    scope_text=None,
    resource_schedule_text=None,
    communication_plan_text=None,
):
    """
    Replace placeholders in the template:
    <<EXEC_SUMMARY>>, <<OBJECTIVE>>, <<SCOPE_TEXT>>, <<RESOURCE_SCHEDULE>>, <<COMMUNICATION_PLAN>>
    Now includes robust bullet point handling.
    """

    doc = Document(template_path)

    def set_cell_shading(cell, fill_color):
        """Add shading (background color) to a table cell."""
        tc_pr = cell._element.tcPr
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tc_pr.append(shd)

    def set_table_border_white(table, cell_margin=150):
        """Set all table borders to white (for clean, minimal look)."""
        tbl = table._element
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement("w:tblBorders")

        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "single")
            border_el.set(qn("w:sz"), "4")  # thin border
            border_el.set(qn("w:space"), "0")
            border_el.set(qn("w:color"), "FFFFFF")  # white
            tbl_borders.append(border_el)

        tbl_pr.append(tbl_borders)


    def insert_styled_table(parent, headers, rows):
        """Create a table styled similar to RFP objective section."""
        table = parent.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = True

        # Header row styling
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.strip()
            set_cell_shading(hdr_cells[i], "008FD3")  # blue header
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for r, row_data in enumerate(rows):
            cells = table.rows[r + 1].cells
            for c, val in enumerate(row_data):
                cells[c].text = str(val).strip()
                set_cell_shading(cells[c], "E7EEF7")  # light gray row
                cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[c].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )

        # Set uniform width
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3)

        # Apply white borders
        set_table_border_white(table)

        return table

    def replace_placeholder(doc, placeholder, new_text):
        if not new_text:
            return

        for para in doc.paragraphs:
            if placeholder in "".join(run.text for run in para.runs):
                parent = para._element.getparent()
                idx = parent.index(para._element)
                parent.remove(para._element)

                lines = [line.strip() for line in new_text.split("\n") if line.strip()]
                new_elements = []  # collect to insert once

                i = 0
                while i < len(lines):
                    line = lines[i]

                    # Markdown-style table
                    if line.startswith("|") and "|" in line:
                        table_lines = []
                        while i < len(lines) and lines[i].startswith("|"):
                            table_lines.append(lines[i])
                            i += 1
                        headers = [h.strip("* ") for h in table_lines[0].strip("|").split("|")]
                        rows = [
                            [c.strip() for c in r.strip("|").split("|")]
                            for r in table_lines[2:]
                        ]
                        table = insert_styled_table(doc, headers, rows)
                        new_elements.append(table._element)
                        continue

                    # Section heading
                    if line.startswith("**") or line.startswith("###"):
                        header_text = line.strip("*# ").rstrip(":")
                        new_para = doc.add_paragraph(header_text)
                        new_para.style = "Table Column Heading"
                        new_para.paragraph_format.space_after = Pt(4)
                        new_elements.append(new_para._element)
                        i += 1
                        continue

                    # Markdown bullets (FIXED: Use apply_bullet_to_para for robustness)
                    if line.startswith("- ") or line.startswith("• "):
                        bullet_text = line[2:].strip() if line.startswith("- ") else line[1:].strip()
                        new_para = doc.add_paragraph(bullet_text, style="List Bullet 2")
                        new_para.paragraph_format.left_indent = Pt(18)
                        new_para.paragraph_format.space_after = Pt(2)
                        new_elements.append(new_para._element)
                        i += 1
                        continue

                    # Regular text
                    new_para = doc.add_paragraph(line)
                    new_elements.append(new_para._element)
                    i += 1

                # ⚡️ Insert all new elements once
                for element in reversed(new_elements):
                    parent.insert(idx, element)
                return

    # Replace placeholders with sections
    replace_placeholder(doc, "<<EXEC_SUMMARY>>", summary_text)
    replace_placeholder(doc, "<<OBJECTIVE>>", objective_text)
    replace_placeholder(doc, "<<SCOPE_TEXT>>", scope_text)
    replace_placeholder(doc, "<<RESOURCE_SCHEDULE>>", resource_schedule_text)
    replace_placeholder(doc, "<<COMMUNICATION_PLAN>>", communication_plan_text)

    return doc


def generate_exec_summary_and_objective(reference_text, condensed_rfp, num_interfaces=113):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

    ref_snippet = get_most_relevant_reference(
        reference_text, condensed_rfp, keyword_hint="executive summary overview introduction proposal background"
    )

    # --- Step 3: Prevent input overflow ---
    if len(ref_snippet) > 8000:
        ref_snippet = ref_snippet[:8000]
    if len(condensed_rfp) > 8000:
        condensed_rfp = condensed_rfp[:8000]


    prompt = get_executive_summary_and_objective_prompt(ref_snippet, condensed_rfp, num_interfaces)

        # --- Prevent input overflow ---
    if len(reference_text) > 8000:
        reference_text = reference_text[:8000]
    if len(condensed_rfp) > 8000:
        condensed_rfp = condensed_rfp[:8000]


    response = client.chat.completions.create(
        model="4o",
        temperature=0.3,
        max_tokens=1800,
        messages=[{"role": "user", "content": prompt}]
    )

    full_output = response.choices[0].message.content.strip()

    # --- Split into Executive Summary and Objective ---
    exec_match = re.search(r"\*\*?Executive Summary\*\*?\s*(.*?)\s*(?=\*\*?Objective\*\*?)", full_output, re.S | re.I)
    obj_match = re.search(r"\*\*?Objective\*\*?\s*(.*)", full_output, re.S | re.I)

    exec_text = exec_match.group(1).strip() if exec_match else full_output
    obj_text = obj_match.group(1).strip() if obj_match else ""

    return exec_text, obj_text

def generate_scope_sections(reference_text, condensed_rfp, num_interfaces=None):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )
        # 🔹 Extract relevant reference for scope section
    ref_snippet = get_most_relevant_reference(
        reference_text, condensed_rfp, keyword_hint="scope assumptions prerequisites out of scope deliverables migration"
    )

    if len(ref_snippet) > 8000:
        ref_snippet = ref_snippet[:8000]
    if len(condensed_rfp) > 8000:
        condensed_rfp = condensed_rfp[:8000]


    prompt = get_scope_prereq_assumptions_prompt(ref_snippet, condensed_rfp, num_interfaces)

    response = client.chat.completions.create(
        model="Codetest",
        temperature=0.3,
        max_tokens=900,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

def generate_resource_schedule_and_commercial(reference_text, condensed_rfp):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

        # 🔹 Extract relevant reference for resource schedule
    ref_snippet = get_most_relevant_reference(
        reference_text, condensed_rfp, keyword_hint="resource schedule commercials cost estimation effort team structure staffing"
    )

    if len(ref_snippet) > 8000:
        ref_snippet = ref_snippet[:8000]
    if len(condensed_rfp) > 8000:
        condensed_rfp = condensed_rfp[:8000]


    prompt = get_resource_schedule_and_commercial_prompt(ref_snippet, condensed_rfp)

    response = client.chat.completions.create(
        model="4o",
        temperature=0.3,
        max_tokens=1800,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()

def generate_communication_plan(reference_text, condensed_rfp):
    # ... (function body remains the same)
    client = AzureOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
        api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
        api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
    )

        # 🔹 Extract relevant reference for communication plan
    ref_snippet = get_most_relevant_reference(
        reference_text, condensed_rfp, keyword_hint="communication governance escalation reporting meetings project coordination issue management"
    )

    if len(ref_snippet) > 8000:
        ref_snippet = ref_snippet[:8000]
    if len(condensed_rfp) > 8000:
        condensed_rfp = condensed_rfp[:8000]


    prompt = get_communication_plan_prompt(ref_snippet, condensed_rfp)
    response = client.chat.completions.create(
        model="gpt-35-turbo",
        temperature=0.3,
        max_tokens=2300,
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content.strip()





# --- Step 1: Upload ---
st.markdown("## 📥 Step 1: Upload Your RFP Document")

uploaded_file = st.file_uploader(
    " ",
    type=["pdf", "docx"],
    key="rfp_uploader",
    help="Upload your RFP document in PDF or DOCX format.",
    label_visibility="collapsed"
)

st.markdown("</div>", unsafe_allow_html=True)

# --- Dynamic Input ---
st.markdown("---")
st.markdown("### ⚙️ Proposal Configuration")


# --- Conditional Logic ---
if uploaded_file:

        st.markdown("### ✍️ Step 2: Generating Your Proposal Response")
        with st.spinner("Analyzing RFP and preparing your AI-driven proposal response..."):


            with st.status("🚀 Generating Proposal Sections...", expanded=True) as status:
        
                # STEP 1: Extract content
                st.write("1/6 🔎 Extracting RFP content...")
                # rfp_text = extract_text(uploaded_file)
                rfp_text = extract_text(uploaded_file)
                st.write("🧠 Condensing extracted RFP content for faster processing...")
                client = AzureOpenAI(
                    azure_endpoint=os.getenv("AZURE_OPENAI_FRFP_ENDPOINT"),
                    api_key=os.getenv("AZURE_OPENAI_FRFP_KEY"),
                    api_version=os.getenv("AZURE_OPENAI_FRFP_VERSION")
                )
                condensed_rfp = condense_rfp_text(rfp_text, client)



                time.sleep(1)
                # --- 🔍 Auto-detect number of interfaces / integrations from RFP text ---
                # import re
                rfp_text = rfp_text.replace(",", "")

                priority_keywords = ["ICOs?", "iCos?", "integration configuration objects?"]
                general_keywords = [
                    "interfaces?", "integration points?", "flows?", "connections?",
                    "touchpoints?", "IFlows?", "mappings?", "adapters?"
                ]

                # First: look specifically for ICO mentions
                ico_pattern = r'~?\b(\d{1,5})\s*(?:' + "|".join(priority_keywords) + r')\b'
                ico_matches = re.findall(ico_pattern, rfp_text, flags=re.IGNORECASE)

                if ico_matches:
                    num_interfaces = max(map(int, ico_matches))
                    detected_type = "ICOs"
                else:
                    # fallback to general terms like 'interfaces' if ICOs not found
                    pattern = r'~?\b(\d{1,5})\s*(?:' + "|".join(general_keywords) + r')\b'
                    matches = re.findall(pattern, rfp_text, flags=re.IGNORECASE)

                    if matches:
                        num_interfaces = max(map(int, matches))
                        detected_type = "interfaces"
                    else:
                        num_interfaces = None
                        detected_type = None

                                # Display result
                if num_interfaces:
                    st.info(f"📊 Detected approximately **{num_interfaces} {detected_type}** in RFP.")
                else:
                    st.warning("⚠️ No explicit integration count detected — using default or manual input.")

                
                if len(rfp_text.strip()) < 100:
                    status.update(label="Extraction Failed", state="error", expanded=False)
                    st.error("Could not extract enough text from the document. Please check the file.")
                    st.stop()
                
                st.success("1/6 ✅ RFP content extracted!")
                status.update(label="🚀 Generating Proposal Sections... (20% Complete)", state="running")

            

                # STEP 2: Build or load knowledge base & Retrieve context
                st.write("2/6 📚 Loading knowledge base and retrieving reference documents...")
                t0 = time.time()
                knowledge_db = build_knowledge_base()
                # elapsed_kb = round(time.time() - t0, 1)
                # st.success(f"✅ Knowledge base loaded in {elapsed_kb} seconds")

                retriever = knowledge_db.as_retriever(search_kwargs={"k": 3})
                ref_docs = retriever.invoke(rfp_text)

                # ✅ Show actual retrieved sources (unique)
                unique_sources = list({d.metadata.get("source", "Unknown") for d in ref_docs})
                if unique_sources:
                    st.success(f"2/6 ✅ Retrieved {len(unique_sources)} unique reference document(s): {', '.join(unique_sources)}")
                else:
                    st.warning("⚠️ No relevant documents retrieved from the knowledge base.")

                reference_text = "\n\n".join([d.page_content[:2500] for d in ref_docs])
                max_chars = 8000
                if len(reference_text) > max_chars:
                    reference_text = reference_text[:max_chars]

                # 🔍 Reduce reference size to only most relevant parts before prompt creation
                reference_text = get_most_relevant_reference(reference_text, condensed_rfp)

                status.update(label="🚀 Generating Proposal Sections... (40% Complete)", state="running")

                # ---------------------------------------------------------------
                # ✅ PARALLEL EXECUTION STARTS HERE
                # ---------------------------------------------------------------
                from concurrent.futures import ThreadPoolExecutor

                st.write("🚀 Generating all proposal sections ")
                status.update(label="🚀 Generating all proposal sections...", state="running")

                start_time = time.time()
                with ThreadPoolExecutor(max_workers=4) as executor:
                    futures = {
                        "exec_obj": executor.submit(generate_exec_summary_and_objective, reference_text, rfp_text, num_interfaces),
                        "scope": executor.submit(generate_scope_sections, reference_text, rfp_text, num_interfaces),
                        "resource": executor.submit(generate_resource_schedule_and_commercial, reference_text, rfp_text),
                        "comm_plan": executor.submit(generate_communication_plan, reference_text, rfp_text),
                    }

                    results = {}
                    for name, future in futures.items():
                        try:
                            results[name] = future.result()
                            st.success(f"✅ {name.replace('_', ' ').title()} generated successfully!")
                        except Exception as e:
                            st.error(f"❌ Failed to generate {name}: {e}")

                # Extract results safely
                exec_summary, objective = results["exec_obj"]
                scope_text = results["scope"]
                resource_schedule_text = results["resource"]
                communication_plan_text = results["comm_plan"]

                end_time = time.time()
                elapsed = round(end_time - start_time, 1)
                st.success(f"✨ All sections generated in {elapsed} seconds!")
                status.update(label="✅ Proposal Content Complete!", state="complete", expanded=False)
            
            
            # --- Proposal Preview ---
            st.markdown("## 🔍 Step 2: Review and Edit Content")
            st.info("Review the AI-generated sections below before downloading the final document.")
            
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "Executive Summary", "Objective", "Scope & Assumptions", 
                "Resource & Schedule", "Communication Plan"
            ])
            
            with tab1: st.markdown(exec_summary)
            with tab2: st.markdown(objective)
            with tab3: st.markdown(scope_text)
            with tab4: st.markdown(resource_schedule_text)
            with tab5: st.markdown(communication_plan_text)
            
            # --- Download Section ---
            st.markdown("---")
            st.markdown("## 📦 Step 3: Final Document Generation & Download")
            template_path = "Template/PIPO TO IS Response Template.docx"

            if not os.path.exists(template_path):
                st.error(f"Template not found at {template_path}. Cannot generate final DOCX.")
            else:
                st.write("Compiling content into DOCX template...")
                final_doc = insert_executive_summary_into_template(
                    template_path,
                    summary_text=exec_summary,
                    objective_text=objective,
                    scope_text=scope_text,
                    resource_schedule_text=resource_schedule_text,
                    communication_plan_text=communication_plan_text
                )

                buffer = BytesIO()
                final_doc.save(buffer)
                buffer.seek(0)
                
                st.markdown("<br>", unsafe_allow_html=True)
                st.download_button(
                    label="🚀 Download Final RFP Proposal (DOCX)",
                    data=buffer,
                    file_name=f"RFP_Response_{uploaded_file.name.split('.')[0]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


            st.success("✅ Proposal response generated successfully!")
            # st.write(result) 
else:
    st.info("⚠️ Please upload an RFP document (PDF or DOCX) before generating.")

